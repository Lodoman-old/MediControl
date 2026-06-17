from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = "screenshots"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Helper functions ──
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    for r in rows:
        row = table.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = str(val)
    return table

def add_screenshot(name, caption, width=Inches(5.5)):
    path = os.path.join(OUT_DIR, f"{name}.png")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(100, 100, 100)

def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def add_bullets(items):
    for item in items:
        add_bullet(item)

def add_field_table(fields):
    """fields: list of (label, description)"""
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, desc) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = desc
    return table

# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MediControl")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0, 82, 136)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Gestión Clínica")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manual de Usuario")
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versión 1.0")
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# INDICE
# ═══════════════════════════════════════════════
add_heading("Índice", level=1)
toc_items = [
    "1. Introducción",
    "2. Acceso al Sistema",
    "  2.1 Pantalla de Login",
    "  2.2 Recuperación de Contraseña",
    "3. Dashboard Principal",
    "4. Agenda de Citas",
    "  4.1 Calendario Semanal",
    "  4.2 Crear Cita",
    "  4.3 Confirmar Cita",
    "  4.4 Estados de Cita",
    "5. Módulo de Enfermería (Triage)",
    "  5.1 Captura de Signos Vitales",
    "  5.2 Transición a Consulta",
    "6. Expediente Clínico",
    "  6.1 Lista de Pacientes",
    "  6.2 Historia Médica",
    "  6.3 Notas de Evolución (SOAP)",
    "  6.4 Diagnósticos (CIE-10)",
    "  6.5 Tratamientos",
    "  6.6 Recetas y PDF",
    "  6.7 Órdenes de Estudio",
    "  6.8 Consentimientos Informados",
    "7. Farmacia",
    "  7.1 Surtir Receta",
    "  7.2 Venta POS",
    "  7.3 Inventario",
    "8. Administración",
    "  8.1 Usuarios y Roles",
    "  8.2 Servicios",
    "  8.3 Ubicaciones",
    "  8.4 Horarios",
    "9. Reportes",
    "10. Solución de Problemas",
    "   10.1 Error 401 / Sesión Expirada",
    "   10.2 La Página no Carga",
    "   10.3 Cold Start en Render",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════
add_heading("1. Introducción", level=1)
add_para("MediControl es un sistema de gestión clínica multi-tenant diseñado para consultorios, clínicas y hospitales. Permite administrar agendas, expedientes clínicos electrónicos (NOM-004-SSA3-2012), farmacia, recetas con PDF y más.")
add_para("Usuarios demo disponibles:")
creds = [
    ("admin@medicontrol.mx", "Admin123!Demo", "Administrador"),
    ("doctor@medicontrol.mx", "Doctor123!Demo", "Médico"),
    ("enfermera@medicontrol.mx", "Enfermera123!Demo", "Enfermera"),
    ("cajero@medicontrol.mx", "Cajero123!Demo", "Cajero"),
    ("paciente@medicontrol.mx", "Paciente123!Demo", "Paciente"),
]
add_table(["Rol", "Email", "Contraseña"], [(r, e, p) for e, p, r in creds])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. ACCESO AL SISTEMA
# ═══════════════════════════════════════════════
add_heading("2. Acceso al Sistema", level=1)

add_heading("2.1 Pantalla de Login", level=2)
add_para("Al abrir la aplicación, se muestra la pantalla de inicio de sesión. En dispositivos móviles tiene un diseño de fondo gradiente; en escritorio muestra un split con el logo.")
add_screenshot("01_login", "Figura 1: Pantalla de inicio de sesión")

add_para("Pasos para iniciar sesión:", bold=True)
add_bullets([
    "Ingresa tu correo electrónico",
    "Ingresa tu contraseña",
    "Haz clic en 'Ingresar'",
    "Si es tu primer inicio, se te pedirá cambiar la contraseña",
])

add_heading("2.2 Recuperación de Contraseña", level=2)
add_para("En la pantalla de login hay un enlace '¿Olvidaste tu contraseña?' que envía un correo de recuperación (funcionalidad disponible si el módulo de email está configurado).")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. DASHBOARD
# ═══════════════════════════════════════════════
add_heading("3. Dashboard Principal", level=1)
add_para("Al iniciar sesión, el dashboard muestra un resumen con tarjetas de:")
add_bullets([
    "Citas de hoy (total, pendientes, completadas)",
    "Pacientes en espera (CHECKED_IN / IN_TRIAGE)",
    "Próximas citas del día",
])
add_screenshot("03_dashboard", "Figura 2: Dashboard principal")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. AGENDA
# ═══════════════════════════════════════════════
add_heading("4. Agenda de Citas", level=1)

add_heading("4.1 Calendario Semanal", level=2)
add_para("La agenda muestra una cuadrícula semanal con los slots disponibles y ocupados. Los colores indican el estado:")
add_field_table([
    ("PENDIENTE", "Amarillo - Pendiente de confirmación del médico"),
    ("CONFIRMADA", "Verde - Cita confirmada"),
    ("EN TRIAGE", "Verde azulado - Paciente siendo atendido por enfermería"),
    ("EN CONSULTA", "Azul - Paciente con el médico"),
    ("COMPLETADA", "Gris oscuro - Consulta finalizada"),
    ("NO SHOW", "Rojo - Paciente no asistió"),
    ("CANCELADA", "Gris claro - Cita cancelada"),
])
add_screenshot("02_agenda", "Figura 3: Agenda semanal")

add_heading("4.2 Crear Cita", level=2)
add_para("Para agendar una nueva cita:")
add_bullets([
    "Haz clic en 'Nueva Cita'",
    "Busca el paciente por nombre (autocomplete)",
    "Selecciona el servicio (determina la duración del slot)",
    "Elige la hora de inicio (el sistema calcula el fin automáticamente)",
    "Selecciona el médico (o se auto-selecciona si eres doctor)",
    "Guarda la cita. Queda en estado PENDING_DOCTOR_CONFIRMATION",
])

add_heading("4.3 Confirmar Cita", level=2)
add_para("El médico debe confirmar la cita y asignar la ubicación (consultorio). Esto cambia el estado a SCHEDULED.")

add_heading("4.4 Estados de Cita", level=2)
add_para("Flujo completo de estados:")
add_bullets([
    "PENDING_DOCTOR_CONFIRMATION → SCHEDULED (médico confirma + asigna ubicación)",
    "SCHEDULED → PAID (si aplica pago) o directo a CHECKED_IN",
    "CHECKED_IN → IN_TRIAGE (enfermería captura signos vitales)",
    "IN_TRIAGE → IN_CONSULT (médico inicia consulta)",
    "IN_CONSULT → COMPLETED (médico finaliza)",
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. ENFERMERIA / TRIAGE
# ═══════════════════════════════════════════════
add_heading("5. Módulo de Enfermería (Triage)", level=1)

add_heading("5.1 Captura de Signos Vitales", level=2)
add_para("La enfermera captura los signos vitales del paciente antes de pasar a consulta. Campos disponibles:")
add_field_table([
    ("PA Sistólica / Diastólica", "Presión arterial (mmHg)"),
    ("Frecuencia Cardíaca", "Latidos por minuto"),
    ("Frecuencia Respiratoria", "Respiraciones por minuto"),
    ("Temperatura", "Temperatura corporal (°C)"),
    ("Saturación de Oxígeno", "SpO2 (%)"),
    ("Peso", "Peso corporal (kg)"),
    ("Talla", "Estatura (m)"),
    ("Glucosa", "Glucosa capilar (mg/dL)"),
    ("Notas", "Observaciones adicionales"),
])
add_para("Al guardar, la cita pasa a estado IN_CONSULT para que el médico inicie la consulta.")

add_heading("5.2 Transición a Consulta", level=2)
add_para("Una vez capturados los signos, el médico puede iniciar la consulta desde la agenda o desde el detalle de la cita.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. EXPEDIENTE
# ═══════════════════════════════════════════════
add_heading("6. Expediente Clínico", level=1)

add_heading("6.1 Lista de Pacientes", level=2)
add_para("Muestra todos los pacientes registrados con búsqueda por nombre, MRN o CURP.")
add_screenshot("04_patient_list", "Figura 4: Lista de pacientes")

add_para("Cada paciente tiene un enlace 'Expediente' que abre su historial clínico completo.")

add_heading("6.2 Historia Médica", level=2)
add_para("La historia médica se compone de:")
add_field_table([
    ("Antecedentes heredofamiliares", "Enfermedades familiares (diabetes, hipertensión, etc.)"),
    ("Antecedentes no patológicos", "Tabaquismo, alcoholismo, actividad física"),
    ("Antecedentes patológicos", "Cirugías previas, enfermedades crónicas, alergias"),
    ("Padecimiento actual", "Motivo de consulta y evolución"),
    ("Revisión por sistemas", "Hallazgos por aparatos y sistemas"),
])
add_para("Todos los campos son texto libre. Anteriormente se usaba formato JSON, ahora es texto natural.")
add_screenshot("06_medical_history", "Figura 5: Formulario de historia médica")

add_heading("6.3 Notas de Evolución (SOAP)", level=2)
add_para("Formato SOAP (Subjetivo, Objetivo, Evaluación, Plan):")
add_field_table([
    ("Subjetivo (S)", "Lo que refiere el paciente (síntomas, molestias)"),
    ("Objetivo (O)", "Hallazgos de exploración física y estudios"),
    ("Evaluación (A)", "Impresión diagnóstica y análisis"),
    ("Plan (P)", "Tratamiento, estudios, seguimiento"),
])
add_screenshot("05_expediente", "Figura 6: Expediente clínico con tabs")

add_heading("6.4 Diagnósticos (CIE-10)", level=2)
add_para("Los diagnósticos se registran con código CIE-10 y pueden ser:")
add_bullets([
    "PRINCIPAL - Diagnóstico principal de la consulta",
    "SECONDARY - Diagnósticos secundarios o comorbilidades",
    "DIFFERENTIAL - Diagnósticos diferenciales por descartar",
])
add_para("Estados: ACTIVO, RESUELTO, SOSPECHADO.")

add_heading("6.5 Tratamientos", level=2)
add_para("Tipos de tratamiento:")
add_bullets([
    "FARMACOLÓGICO - Medicamentos con dosis, frecuencia y duración",
    "NO FARMACOLÓGICO - Fisioterapia, terapia, cambios de estilo de vida",
    "QUIRÚRGICO - Procedimientos quirúrgicos",
])
add_para("Estados: ACTIVO, COMPLETADO, CANCELADO, EN ESPERA.")

add_heading("6.6 Recetas y PDF", level=2)
add_para("Las recetas médicas se generan en formato PDF con:")
add_bullets([
    "Encabezado: DR. [nombre], especialidad, cédulas profesional y de especialidad en el mismo renglón",
    "Datos del paciente: nombre, edad, fecha",
    "Medicamento: nombre, dosis, frecuencia, duración, vía de administración",
    "Instrucciones e indicaciones",
    "Signos vitales de la consulta (opcional)",
    "Próxima cita programada",
    "Línea de firma autógrafa o digital",
    "Formato NOM-072-SSA3-2012",
])
add_para("El PDF se descarga autenticado (requiere sesión activa). Hay botón 'PDF' tanto en el expediente como en el detalle de cita.")

add_heading("6.7 Órdenes de Estudio", level=2)
add_para("El médico puede solicitar estudios de laboratorio, imagen, patología u otros. Cada orden tiene:")
add_bullets([
    "Tipo de estudio (Laboratorio, Imagen, Patología, Otro)",
    "Nombre del estudio específico",
    "Indicación o justificación clínica",
    "Estado: Pendiente → Tomado → En proceso → Completado",
])
add_para("También se genera PDF para la orden de estudio.")

add_heading("6.8 Consentimientos Informados", level=2)
add_para("Tipos de consentimiento:")
add_field_table([
    ("GENERAL", "Consentimiento general para atención médica"),
    ("CIRUGÍA", "Procedimientos quirúrgicos"),
    ("ANESTESIA", "Administración de anestesia"),
    ("TRANSFUSIÓN", "Transfusión de sangre o hemoderivados"),
    ("INVESTIGACIÓN", "Participación en estudios de investigación"),
    ("OTROS", "Otros procedimientos específicos"),
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 7. FARMACIA
# ═══════════════════════════════════════════════
add_heading("7. Farmacia", level=1)

add_heading("7.1 Surtir Receta", level=2)
add_para("El módulo de farmacia permite surtir (dispensar) recetas activas. Al surtir:")
add_bullets([
    "Se reduce el inventario automáticamente",
    "Se registra la dispensación (quién, cuándo, cantidad)",
    "La receta cambia a estado DISPENSED o PARTIALLY_DISPENSED",
])

add_heading("7.2 Venta POS", level=2)
add_para("El módulo POS permite vender productos de farmacia sin receta. Incluye:")
add_bullets([
    "Selección de sucursal, método de pago y paciente",
    "Búsqueda de productos por código o nombre",
    "Carrito de compra con cantidades",
    "Ticket de venta",
])

add_heading("7.3 Inventario", level=2)
add_para("Gestiona el stock de medicamentos y productos. Permite movimientos de entrada, salida y ajuste.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 8. ADMINISTRACIÓN
# ═══════════════════════════════════════════════
add_heading("8. Administración", level=1)

add_heading("8.1 Usuarios y Roles", level=2)
add_para("Solo SUPERADMIN y ADMIN pueden gestionar usuarios. Roles disponibles:")
add_bullets([
    "SUPERADMIN - Acceso total al sistema",
    "ADMIN - Gestión administrativa",
    "DOCTOR - Atención médica, expediente, recetas",
    "NURSE - Captura de signos vitales (triage)",
    "RECEPTION - Gestión de citas y pacientes",
    "CAJERO - Cobros y farmacia POS",
    "PATIENT - Auto-consulta de citas (portal paciente)",
])

add_heading("8.2 Servicios", level=2)
add_para("Los servicios definen el tipo de consulta y su duración. Cada servicio tiene: código, nombre, precio y duración en minutos. La duración determina el tamaño del slot en la agenda.")

add_heading("8.3 Ubicaciones", level=2)
add_para("Las ubicaciones (consultorios, salas) se asignan al confirmar la cita. Tipos: EXAM_ROOM, OR, BED, ER_BAY, PROCEDURE, LAB, IMAGING.")

add_heading("8.4 Horarios", level=2)
add_para("Los horarios del médico se configuran en una cuadrícula semanal con modal CRUD. Soporta múltiples bloques por día sin superposición. El horario se usa para calcular slots disponibles.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 9. REPORTES
# ═══════════════════════════════════════════════
add_heading("9. Reportes", level=1)
add_para("El módulo de reportes incluye:")
add_bullets([
    "Reporte de ventas (historial y resumen)",
    "Corte de caja (apertura/cierre con conteo de efectivo)",
    "Reporte de citas (por médico, servicio, período)",
    "Reporte de pacientes atendidos",
])
add_para("Las tablas tienen scroll horizontal para mejor visualización en móvil.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# 10. SOLUCIÓN DE PROBLEMAS
# ═══════════════════════════════════════════════
add_heading("10. Solución de Problemas", level=1)

add_heading("10.1 Error 401 / Sesión Expirada", level=2)
add_para("Síntoma: La aplicación muestra error 401 y redirige al login.")
add_para("Causa: El token JWT ha expirado (15 min de acceso, 7 días de refresh). O las claves JWT cambiaron en un nuevo deploy.")
add_para("Solución:")
add_bullets([
    "Abre el sitio en una ventana de incógnito",
    "O limpia el sessionStorage desde DevTools > Application > Session Storage > Clear All",
    "Inicia sesión nuevamente",
])

add_heading("10.2 La Página no Carga / 404", level=2)
add_para("Síntoma: Al recargar o navegar directamente a una ruta (ej: /agenda) aparece 404.")
add_para("Causa: El servidor estático no tiene configuración SPA (single page application).")
add_para("Solución en Render: Ir a Dashboard > medicontrol-web > Redirects/Rewrites > agregar regla: Source: /*, Destination: /index.html, Action: Rewrite. O sincronizar el Blueprint.")

add_heading("10.3 Cold Start en Render", level=2)
add_para("Síntoma: La primera carga del día tarda 30-40 segundos.")
add_para("Causa: Render free tier duerme el servicio después de 15 min de inactividad.")
add_para("Solución: Esperar a que cargue. No recargar durante el cold start. Las peticiones posteriores serán rápidas.")

# ── FOOTER ──
doc.add_page_break()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fin del Manual —")
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MediControl © 2026 — Versión 1.0")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(130, 130, 130)

OUT_PATH = "MANUAL_DE_USUARIO.docx"
doc.save(OUT_PATH)
print(f"Manual generado: {OUT_PATH}")
